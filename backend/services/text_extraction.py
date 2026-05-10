"""文档文本提取工具。

统一为 Wiki 生成、原文检索和文档预览提供文本提取能力。
"""

import re
from pathlib import Path
from typing import Dict, List


SUPPORTED_TEXT_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt", ".md", ".markdown"}


def _clean_text(text: str) -> str:
    """清理提取后的文本。"""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f]", "", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _read_text_file(file_path: Path) -> str:
    """读取纯文本 / Markdown 文件，带编码兜底。"""
    raw = file_path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return _clean_text(raw.decode(encoding))
        except UnicodeDecodeError:
            continue
    return _clean_text(raw.decode("latin-1", errors="ignore"))


def _read_text_pages(file_path: Path) -> List[Dict[str, str | int]]:
    """读取纯文本 / Markdown 文件，作为单页返回。"""
    text = _read_text_file(file_path)
    return [{"page": 1, "text": text}]


def _read_docx_text(file_path: Path) -> str:
    """提取 DOCX 文本，包含段落和表格。"""
    from docx import Document

    doc = Document(file_path)
    parts: List[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return _clean_text("\n".join(parts))


def _read_docx_pages(file_path: Path) -> List[Dict[str, str | int]]:
    """提取 DOCX 文本，作为单页返回。"""
    return [{"page": 1, "text": _read_docx_text(file_path)}]


def _decode_ansi_piece(raw: bytes) -> str:
    """对 DOC 的压缩文本片段做多编码兜底解码。"""
    candidates = ("gb18030", "gbk", "cp936", "cp1252", "latin1")
    best_text = ""
    best_score = -1

    for encoding in candidates:
        try:
            text = raw.decode(encoding, errors="strict")
        except UnicodeDecodeError:
            continue

        score = sum(ch.isprintable() or ch in "\n\t" for ch in text)
        score += sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff") * 2
        if score > best_score:
            best_text = text
            best_score = score

    if best_text:
        return best_text
    return raw.decode("latin1", errors="ignore")


def _find_piece_table(clx: bytes) -> bytes | None:
    """从 Clx 中定位 piece table。"""
    pos = 0
    while pos < len(clx):
        marker = clx[pos]
        if marker == 0x02 and pos + 5 <= len(clx):
            lcb_piece_table = int.from_bytes(clx[pos + 1:pos + 5], "little")
            start = pos + 5
            end = min(len(clx), start + lcb_piece_table)
            return clx[start:end]
        if marker == 0x01 and pos + 2 <= len(clx):
            pos += 2 + clx[pos + 1]
            continue
        break
    return None


def _read_doc_text(file_path: Path) -> str:
    """提取 DOC 文本，基于 MS-DOC piece table。"""
    import olefile

    if not olefile.isOleFile(file_path):
        return "[DOC解析错误: 文件不是有效的 OLE 复合文档]"

    with olefile.OleFileIO(file_path) as ole:
        if not ole.exists("WordDocument"):
            return "[DOC解析错误: 缺少 WordDocument 流]"

        word_stream = ole.openstream("WordDocument").read()
        fib = word_stream[:1472]
        if len(fib) < 0x1AA:
            return "[DOC解析错误: FIB 结构不完整]"

        flag_word = int.from_bytes(fib[0x000A:0x000E], "little")
        table_stream_name = "1Table" if (flag_word & 0x0200) else "0Table"
        if not ole.exists(table_stream_name):
            return f"[DOC解析错误: 缺少 {table_stream_name} 流]"

        table_stream = ole.openstream(table_stream_name).read()
        fc_clx = int.from_bytes(fib[0x01A2:0x01A6], "little")
        lcb_clx = int.from_bytes(fib[0x01A6:0x01AA], "little")
        if fc_clx >= len(table_stream) or lcb_clx <= 0:
            return "[DOC解析错误: 无法定位 Clx 结构]"

        clx = table_stream[fc_clx:fc_clx + lcb_clx]
        piece_table = _find_piece_table(clx)
        if not piece_table or len(piece_table) < 16:
            return "[DOC解析错误: 无法定位 piece table]"

        piece_count = (len(piece_table) - 4) // 12
        if piece_count <= 0:
            return "[DOC解析错误: piece table 为空]"

        pieces: List[str] = []
        for idx in range(piece_count):
            cp_start = int.from_bytes(piece_table[idx * 4:idx * 4 + 4], "little")
            cp_end = int.from_bytes(piece_table[(idx + 1) * 4:(idx + 1) * 4 + 4], "little")
            if cp_end <= cp_start:
                continue

            descriptor_offset = ((piece_count + 1) * 4) + (idx * 8)
            piece_descriptor = piece_table[descriptor_offset:descriptor_offset + 8]
            if len(piece_descriptor) < 8:
                continue

            fc_value = int.from_bytes(piece_descriptor[2:6], "little")
            is_ansi = (fc_value & 0x40000000) != 0
            fc = fc_value & 0x3FFFFFFF
            char_count = cp_end - cp_start

            if is_ansi:
                start = fc // 2
                raw = word_stream[start:start + char_count]
                pieces.append(_decode_ansi_piece(raw))
            else:
                raw = word_stream[fc:fc + (char_count * 2)]
                pieces.append(raw.decode("utf-16le", errors="ignore"))

    text = _clean_text("\n".join(pieces))
    if not text:
        return "[DOC解析错误: 未提取到可读文本，请确认文件未损坏]"
    return text


def _read_doc_pages(file_path: Path) -> List[Dict[str, str | int]]:
    """提取 DOC 文本，作为单页返回。"""
    return [{"page": 1, "text": _read_doc_text(file_path)}]


def _read_pdf_pages(file_path: Path) -> List[Dict[str, str | int]]:
    """提取 PDF 每页文本。"""
    try:
        import fitz  # PyMuPDF

        pages: List[Dict[str, str | int]] = []
        with fitz.open(file_path) as doc:
            for i, page in enumerate(doc):
                text = _clean_text(page.get_text())
                if text:
                    pages.append({
                        "page": i + 1,
                        "text": text,
                    })

        if not pages:
            return [{"page": 1, "text": "[PDF扫描版: 文档未检测到可提取的文字内容，请上传文字版PDF]"}]
        return pages
    except Exception as e:
        return [{"page": 1, "text": f"[PDF解析错误: {str(e)}]"}]


def extract_document_text(file_path: Path) -> str:
    """按文件类型提取文本内容。"""
    pages = extract_document_pages(file_path)
    return "\n\n".join(page["text"] for page in pages)


def extract_document_pages(file_path: Path) -> List[Dict[str, str | int]]:
    """按文件类型提取分页文本。"""
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _read_pdf_pages(file_path)

    if suffix == ".docx":
        try:
            return _read_docx_pages(file_path)
        except Exception as e:
            return [{"page": 1, "text": f"[Word解析错误: {str(e)}]"}]

    if suffix == ".doc":
        try:
            return _read_doc_pages(file_path)
        except Exception as e:
            return [{"page": 1, "text": f"[Word解析错误: {str(e)}]"}]

    if suffix in [".txt", ".md", ".markdown"]:
        try:
            return _read_text_pages(file_path)
        except Exception as e:
            return [{"page": 1, "text": f"[文本读取错误: {str(e)}]"}]

    return [{"page": 1, "text": f"[不支持的文件格式: {suffix}]"}]
