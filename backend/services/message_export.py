"""消息导出工具。"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


HEADING_RE = re.compile(r"^(\s{0,3})(#{1,6})\s+(.*)$")
ORDERED_LIST_RE = re.compile(r"^(\s*)(\d+)[.)]\s+(.*)$")
UNORDERED_LIST_RE = re.compile(r"^(\s*)([-*+])\s+(.*)$")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")
BR_TAG_RE = re.compile(r"<br\s*\/?>", re.IGNORECASE)


def strip_inline_markdown(text: str) -> str:
    return (
        text.replace("\r\n", "\n")
        .replace("<br>", "\n")
        .replace("<br/>", "\n")
        .replace("<br />", "\n")
        .replace("**", "")
        .replace("__", "")
        .replace("~~", "")
        .replace("`", "")
        .replace("[[", "")
        .replace("]]", "")
        .strip()
    )


def to_plain_text_inline(text: str) -> str:
    cleaned = BR_TAG_RE.sub("\n", text)
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"__([^_]+)__", r"\1", cleaned)
    cleaned = re.sub(r"~~([^~]+)~~", r"\1", cleaned)
    cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
    cleaned = re.sub(r"\[\[([^\]]+)\]\]", r"\1", cleaned)
    cleaned = re.sub(r"\(\s*来源:\s*([^)]+)\s*\)", r"来源：\1", cleaned)
    cleaned = re.sub(r"\[([^\]]+)\]\((https?:\/\/[^\s)]+)\)", r"\1 (\2)", cleaned)
    return cleaned.strip()


def split_table_cells(line: str) -> list[str]:
    return [
        strip_inline_markdown(cell.strip())
        for cell in line.strip().lstrip("|").rstrip("|").split("|")
    ]


def is_table_row(line: str) -> bool:
    return "|" in line.strip()


def is_table_separator(line: str) -> bool:
    return bool(TABLE_SEPARATOR_RE.match(line))


def parse_markdown_blocks(text: str) -> list[dict[str, Any]]:
    lines = text.replace("\r\n", "\n").split("\n")
    blocks: list[dict[str, Any]] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        blocks.append({
            "type": "paragraph",
            "text": strip_inline_markdown(" ".join(paragraph_lines).replace("\n", " ").strip()),
        })
        paragraph_lines = []

    i = 0
    while i < len(lines):
        line = lines[i]
        trimmed = line.strip()

        if not trimmed:
            flush_paragraph()
            i += 1
            continue

        if re.match(r"^(-{3,}|_{3,}|\*{3,})$", trimmed):
            flush_paragraph()
            i += 1
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            flush_paragraph()
            blocks.append({
                "type": "heading",
                "level": len(heading_match.group(2)),
                "text": strip_inline_markdown(heading_match.group(3).strip()),
            })
            i += 1
            continue

        if is_table_row(line) and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            flush_paragraph()
            headers = split_table_cells(line)
            rows: list[list[str]] = []
            cursor = i + 2
            while cursor < len(lines) and is_table_row(lines[cursor]) and lines[cursor].strip():
                rows.append(split_table_cells(lines[cursor]))
                cursor += 1
            blocks.append({
                "type": "table",
                "headers": headers,
                "rows": rows,
            })
            i = cursor
            continue

        ordered_match = ORDERED_LIST_RE.match(line)
        if ordered_match:
            flush_paragraph()
            blocks.append({
                "type": "list-item",
                "ordered": True,
                "indent": len(ordered_match.group(1)) // 2,
                "number": ordered_match.group(2),
                "text": strip_inline_markdown(ordered_match.group(3).strip()),
            })
            i += 1
            continue

        unordered_match = UNORDERED_LIST_RE.match(line)
        if unordered_match:
            flush_paragraph()
            blocks.append({
                "type": "list-item",
                "ordered": False,
                "indent": len(unordered_match.group(1)) // 2,
                "text": strip_inline_markdown(unordered_match.group(3).strip()),
            })
            i += 1
            continue

        paragraph_lines.append(line.strip())
        i += 1

    flush_paragraph()
    return blocks


def _set_doc_font(document: Document, font_name: str = "Microsoft YaHei") -> None:
    styles = document.styles
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        try:
            style = styles[style_name]
        except KeyError:
            continue
        font = style.font
        font.name = font_name
        font.size = font.size or Pt(11)


def build_message_docx_bytes(title: str, content: str) -> bytes:
    document = Document()
    document.core_properties.title = title
    _set_doc_font(document)
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    blocks = parse_markdown_blocks(content)
    if not blocks:
        document.add_paragraph(to_plain_text_inline(content) if content.strip() else "")
    else:
        for block in blocks:
            if block["type"] == "heading":
                level = min(int(block["level"]), 4)
                paragraph = document.add_heading(block["text"], level=level)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue

            if block["type"] == "paragraph":
                paragraph = document.add_paragraph(to_plain_text_inline(block["text"]))
                paragraph.paragraph_format.space_after = Pt(6)
                continue

            if block["type"] == "list-item":
                style = "List Number" if block["ordered"] else "List Bullet"
                paragraph = document.add_paragraph(style=style)
                indent = float(block["indent"]) * 0.25
                paragraph.paragraph_format.left_indent = Inches(indent)
                paragraph.paragraph_format.space_after = Pt(4)
                run = paragraph.add_run(to_plain_text_inline(block["text"]))
                run.font.name = "Microsoft YaHei"
                run.font.size = Pt(10.5)
                continue

            if block["type"] == "table":
                headers = [to_plain_text_inline(header) for header in block["headers"]]
                rows = [[to_plain_text_inline(cell) for cell in row] for row in block["rows"]]
                row_count = max(1, len(rows) + 1)
                col_count = max(1, max([len(headers), *[len(row) for row in rows]] if rows else [len(headers)]))
                table = document.add_table(rows=row_count, cols=col_count)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                if headers:
                    for idx, header in enumerate(headers):
                        if idx >= len(table.rows[0].cells):
                            break
                        table.rows[0].cells[idx].text = header
                for row_index, row in enumerate(rows, start=1):
                    for col_index, cell_text in enumerate(row):
                        if col_index >= len(table.rows[row_index].cells):
                            break
                        table.rows[row_index].cells[col_index].text = cell_text
                continue

    output = BytesIO()
    document.save(output)
    return output.getvalue()
