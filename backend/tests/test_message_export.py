from __future__ import annotations

from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from backend.app import app


def test_export_chat_message_docx_preserves_structured_format():
    client = TestClient(app)

    response = client.post(
        "/api/chat/export-docx",
        json={
            "title": "导出测试",
            "content": "# 一、核心主题\n\n- 先做A\n- 再做B\n\n| 项目 | 要求 |\n| --- | --- |\n| 时间 | 本周 |\n| 责任 | 部门负责人 |",
        },
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    document = Document(BytesIO(response.content))
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    assert any("一、核心主题" in text for text in paragraph_texts)
    assert any("先做A" in text for text in paragraph_texts)
    assert any("再做B" in text for text in paragraph_texts)
    assert len(document.tables) == 1
    table = document.tables[0]
    assert table.cell(0, 0).text == "项目"
    assert table.cell(1, 0).text == "时间"
    assert table.cell(1, 1).text == "本周"
